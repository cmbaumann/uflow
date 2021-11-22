from flask import Flask, render_template, request, url_for, redirect, session, send_file, request
import enum
import re
import pymongo
import bcrypt
from docx import Document
import math
from datetime import date

app = Flask(__name__)
app.secret_key = "testing"
client = pymongo.MongoClient("mongodb+srv://codeusername:CodePassword@testcluster.sl9ku.mongodb.net/test?ssl=true&ssl_cert_reqs=CERT_NONE")

db = client.get_database('total_records')
records = db.register

majors = ['Aerospace Engineering', 'Architecural Engineering', 'Chemical Engineering', 'Civil Engineering',
         'Computer Engineering', 'Computer Science', 'Construction Engineering', 'Cyber Security', 
         'Electircal Engineering', 'Environmental Engineering', 'Mechanical Engineering', 'Metallurgical Engineering',
         'Musical Audio Engineering']

years = [2017, 2018, 2019, 2020, 2021, 2022, 2023]

@app.route('/', methods=['post', 'get'])
def index():
    if request.method == "POST":
        return redirect(url_for('register'))
    return render_template('index.html')

@app.route("/register", methods=['post', 'get'])
def register():
    global majors 
    majors.reverse()
    global years
    years.reverse()
    message = ''
    if "email" in session:
        return redirect(url_for("logged_in"))
    if request.method == "POST":
        first = request.form.get("firstname")
        last = request.form.get("lastname")
        email = request.form.get("email")
        
        password1 = request.form.get("password1")
        password2 = request.form.get("password2")

        major = request.form.get("major")
        year = request.form.get("year")
        
        email_found = records.find_one({"email": email})
        
        if first == "":
            message = 'Please enter your first name'
            return render_template('register.html', message=message, majors=majors, years=years, last=last, email=email, password1=password1, password2=password2, selectedmajor=major, selectedyear=year)
        elif last == "":
            message = 'Please enter your last name'
            return render_template('register.html', message=message, majors=majors, years=years, first=first, email=email, password1=password1, password2=password2, selectedmajor=major, selectedyear=year)
        elif email == "":
            message = 'Please enter your email'
            return render_template('register.html', message=message, majors=majors, years=years, first=first, last=last, password1=password1, password2=password2, selectedmajor=major, selectedyear=year)
        elif (password1 == "") or (password2 == ""):
            message = 'Passwords do not match'
            return render_template('register.html', message=message, majors=majors, years=years, first=first, last=last, email=email, selectedmajor=major, selectedyear=year)
        elif email_found:
            message = 'There is already an account associated with this email'
            return render_template('register.html', message=message, majors=majors, years=years, first=first, last=last, password1=password1, password2=password2, selectedmajor=major, selectedyear=year)
        elif password1 != password2:
            message = 'Passwords do not match'
            return render_template('register.html', message=message, majors=majors, years=years, first=first, last=last, email=email, selectedmajor=major, selectedyear=year)
        else:
            hashed = bcrypt.hashpw(password2.encode('utf-8'), bcrypt.gensalt())
            user_input = {'firstname': first, 'lastname': last, 'email': email, 'password': hashed, 'major': major, 'year': year, 'flowcharts': []}
            records.insert_one(user_input)
            
            user_data = records.find_one({"email": email})
            new_email = user_data['email']
            session["email"] = new_email
   
            return render_template('logged_in.html', email=new_email)
    return render_template('register.html', majors=majors, years=years)

@app.route('/logged_in', methods=['post', 'get'])
def logged_in():
    if request.method == "POST":
        name = request.form.get("name")
        email = session["email"]

        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        newData = []
        len = 0
        for item in entry:
            for thing in item['flowcharts']:
                len += 1
            for i in range(len):
                if (item['flowcharts'][i]["name"] != name):
                    newData.append(item['flowcharts'][i])
        records.update_one(
            {"email": email},
            {"$set": {"flowcharts": newData}}
        )

        names = []
        fcBgColors = []
        data = []
        len = 0
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        for item in entry:
                for thing in item['flowcharts']:
                    len = len + 1
                for i in range(0, len):
                    names.append(item['flowcharts'][i]["name"])
                    fcBgColors.append(item['flowcharts'][i]["colors"][12])
        for item in newData: #each flowchart
            tempData = []
            tempData.append(item["1"])
            tempData.append(item["2"])
            tempData.append(item["3"])
            tempData.append(item["4"])
            tempData.append(item["5"])
            tempData.append(item["6"])
            tempData.append(item["7"])
            tempData.append(item["8"])
            tempData.append(item["9"])
            tempData.append(item["10"])
            tempData.append(item["11"])
            tempData.append(item["12"])
            tempData.append(item["13"])
            tempData.append(item["14"])
            tempData.append(item["15"])
            tempData.append(item["16"])
            tempData.append(item["17"])
            tempData.append(item["18"])
            tempData.append(item["19"])
            tempData.append(item["20"])
            tempData.append(item["21"])
            tempData.append(item["22"])
            tempData.append(item["23"])
            tempData.append(item["24"])
            tempData.append(item["25"])
            tempData.append(item["26"])
            tempData.append(item["27"])
            tempData.append(item["28"])
            tempData.append(item["29"])
            tempData.append(item["30"])
            tempData.append(item["31"])
            tempData.append(item["32"])
            tempData.append(item["33"])
            tempData.append(item["34"])
            tempData.append(item["35"])
            tempData.append(item["36"])
            tempData.append(item["37"])
            data.append(tempData)
            tempData = []
        return render_template('logged_in.html', email=email, names=names, data=data, fcBgColors=fcBgColors)
    if "email" in session:
        email = session["email"]
        names = []
        fcBgColors = []
        data = []
        len = 0
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        for item in entry:
                for thing in item['flowcharts']:
                    len = len + 1
                for i in range(0, len):
                    names.append(item['flowcharts'][i]["name"])
                    fcBgColors.append(item['flowcharts'][i]["colors"][12])
        len = 0
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        print(entry)
        for item in entry: #each flowchart
            for thing in item['flowcharts']:
                len += 1
            for i in range(len): #going through the flowcharts for the user
                    tempData = []
                    tempData.append(item['flowcharts'][i]["1"])
                    tempData.append(item['flowcharts'][i]["2"])
                    tempData.append(item['flowcharts'][i]["3"])
                    tempData.append(item['flowcharts'][i]["4"])
                    tempData.append(item['flowcharts'][i]["5"])
                    tempData.append(item['flowcharts'][i]["6"])
                    tempData.append(item['flowcharts'][i]["7"])
                    tempData.append(item['flowcharts'][i]["8"])
                    tempData.append(item['flowcharts'][i]["9"])
                    tempData.append(item['flowcharts'][i]["10"])
                    tempData.append(item['flowcharts'][i]["11"])
                    tempData.append(item['flowcharts'][i]["12"])
                    tempData.append(item['flowcharts'][i]["13"])
                    tempData.append(item['flowcharts'][i]["14"])
                    tempData.append(item['flowcharts'][i]["15"])
                    tempData.append(item['flowcharts'][i]["16"])
                    tempData.append(item['flowcharts'][i]["17"])
                    tempData.append(item['flowcharts'][i]["18"])
                    tempData.append(item['flowcharts'][i]["19"])
                    tempData.append(item['flowcharts'][i]["20"])
                    tempData.append(item['flowcharts'][i]["21"])
                    tempData.append(item['flowcharts'][i]["22"])
                    tempData.append(item['flowcharts'][i]["23"])
                    tempData.append(item['flowcharts'][i]["24"])
                    tempData.append(item['flowcharts'][i]["25"])
                    tempData.append(item['flowcharts'][i]["26"])
                    tempData.append(item['flowcharts'][i]["27"])
                    tempData.append(item['flowcharts'][i]["28"])
                    tempData.append(item['flowcharts'][i]["29"])
                    tempData.append(item['flowcharts'][i]["30"])
                    tempData.append(item['flowcharts'][i]["31"])
                    tempData.append(item['flowcharts'][i]["32"])
                    tempData.append(item['flowcharts'][i]["33"])
                    tempData.append(item['flowcharts'][i]["34"])
                    tempData.append(item['flowcharts'][i]["35"])
                    tempData.append(item['flowcharts'][i]["36"])
                    tempData.append(item['flowcharts'][i]["37"])
                    data.append(tempData)
                    tempData = []
        return render_template('logged_in.html', email=email, names=names, data=data, fcBgColors=fcBgColors)
    else:
        return redirect(url_for("login"))

@app.route("/login", methods=["POST", "GET"])
def login():
    message = 'Please login to your account'
    if "email" in session:
        return redirect(url_for("logged_in"))

    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")

       
        email_found = records.find_one({"email": email})
        if email_found:
            email_val = email_found['email']
            passwordcheck = email_found['password']
            
            if bcrypt.checkpw(password.encode('utf-8'), passwordcheck):
                session["email"] = email_val
                return redirect(url_for('logged_in'))
            else:
                if "email" in session:
                    return redirect(url_for("logged_in"))
                message = 'Wrong password'
                return render_template('login.html', message=message)
        else:
            message = 'Email not found'
            return render_template('login.html', message=message)
    return render_template('login.html', message=message)

@app.route("/logout", methods=["POST", "GET"])
def logout():
    if "email" in session:
        session.pop("email", None)

        return render_template("signout.html")
    else:
        return render_template('index.html')

def exportData(years, email, name, electives, hours, hoursTaken, hoursPlanned):
    #get student major
    entry = records.find({"email": email}, {"major": 1, "_id": 0})
    for item in entry:
        major = item['major']

    #create document
    document = Document()
    title = major + " Flowchart " + name

    #helper function for determining table size
    def number_helper(value):
        if (value == "taken"): return 1
        elif (value == "inprogress"): return 2
        elif (value == "spring0"): return 3
        elif (value == "fall0"): return 4
        elif (value == "spring1"): return 5
        elif (value == "fall1"): return 6
        elif (value == "spring2"): return 7
        elif (value == "fall2"): return 8
        elif (value == "spring3"): return 9
        elif (value == "fall3"): return 10
        elif (value == "spring4"): return 11
        elif (value == "fall4"): return 12
        elif (value == "deselect"): return 13
        elif (value == ""): return 13

    #helper function to get semester name
    def getSemesterName(num):
        if (num == 1): return "Taken"
        elif (num == 2): return "In Progress"
        elif (num == 3): return "Spring " + str(years[0])
        elif (num == 4): return "Fall " + str(years[0])
        elif (num == 5): return "Spring " + str(years[1])
        elif (num == 6): return "Fall " + str(years[1])
        elif (num == 7): return "Spring " + str(years[2])
        elif (num == 8): return "Fall " + str(years[2])
        elif (num == 9): return "Spring " + str(years[3])
        elif (num == 10): return "Fall " + str(years[3])
        elif (num == 11): return "Spring " + str(years[4])
        elif (num == 12): return "Fall " + str(years[4])
        elif (num == 13): return ""

    def getElectiveName(index):
        if (electives[index] != ""):
            return electives[index] + " (" + str(hours[index]) + ")"
        elif (index == 0): return "HI/SB Elective"
        elif (index == 1): return "HU/L/FA Elective"
        elif (index == 2): return "HU/L/FA Elective"
        elif (index == 3): return "Natural Science Elective"
        elif (index == 4): return "Free Elective"
        elif (index == 5): return "HI/SB Elective"
        elif (index == 6): return "Free Elective"
        elif (index == 7): return "HI/SB Elective"
        elif (index == 8): return "CS 4xx"
        elif (index == 9):  return "Free Elective"
        elif (index == 10): return "HU/L/FA Elective"
        elif (index == 11): return "Natural Sciene Sequence #1"
        elif (index == 12): return "CS 4xx"
        elif (index == 13): return "Free Elective"
        elif (index == 14): return "Free Elective"
        elif (index == 15): return "Natural Science Sequence #2"

    #helper function to get class name based on id
    def getClassName(id):
        if (id == 1): return "EN 101 (3)"
        elif (id == 2): return "ENGR 103(3)"
        elif (id == 3): return "MATH 125 (4)"
        elif (id == 4): return "CS 100 (4)"
        elif (id == 5): return "CS 121 (1)"
        elif (id == 6): return "EN 102 (3)"
        elif (id == 7): return getElectiveName(0)
        elif (id == 8): return "MATH 126 (4)"
        elif (id == 9): return "CS 101 (4)"
        elif (id == 10): return getElectiveName(1)
        elif (id == 11): return "MATH 301 (3)"
        elif (id == 12): return "CS 200 (4)"
        elif (id == 13): return "ECE 380 (4)"
        elif (id == 14): return getElectiveName(2)
        elif (id == 15): return getElectiveName(3)
        elif (id == 16): return "MATH 302 (1)"
        elif (id == 17): return "CS 201 (4)"
        elif (id == 18): return "ECE 383 (4)"
        elif (id == 19): return getElectiveName(4)
        elif (id == 20): return getElectiveName(5)
        elif (id == 21): return "GES 255/ MATH 355 (3)"
        elif (id == 22): return "CS 300(3)"
        elif (id == 23): return "CS 301 (3)"
        elif (id == 24): return getElectiveName(6)
        elif (id == 25): return getElectiveName(7)
        elif (id == 26): return "MATH 237 (3)"
        elif (id == 27): return "CS 403 (3)"
        elif (id == 28): return getElectiveName(8)
        elif (id == 29):  return getElectiveName(9)
        elif (id == 30): return getElectiveName(10)
        elif (id == 31): return getElectiveName(11)
        elif (id == 32): return "CS 470/ CS 475 (3)"
        elif (id == 33): return getElectiveName(12)
        elif (id == 34): return getElectiveName(13)
        elif (id == 35): return getElectiveName(14)
        elif (id == 36): return getElectiveName(15)
        elif (id == 37): return "CS 495 (3)"

    def getElectiveHours(index):
        if (electives[index] != ""):
            return int(hours[index])
        elif (index == 0): return 3
        elif (index == 1): return 3
        elif (index == 2): return 3
        elif (index == 3): return 4
        elif (index == 4): return 3
        elif (index == 5): return 3
        elif (index == 6): return 3
        elif (index == 7): return 3
        elif (index == 8): return 3
        elif (index == 9):  return 3
        elif (index == 10): return 3
        elif (index == 11): return 4
        elif (index == 12): return 3
        elif (index == 13): return 3
        elif (index == 14): return 3
        elif (index == 15): return 4

    def getHours(id):
        if (id == 1): return 3
        elif (id == 2): return 3
        elif (id == 3): return 4
        elif (id == 4): return 4
        elif (id == 5): return 1
        elif (id == 6): return 3
        elif (id == 7): return getElectiveHours(0)
        elif (id == 8): return 4
        elif (id == 9): return 4
        elif (id == 10): return getElectiveHours(1)
        elif (id == 11): return 3
        elif (id == 12): return 4
        elif (id == 13): return 4
        elif (id == 14): return getElectiveHours(2)
        elif (id == 15): return getElectiveHours(3)
        elif (id == 16): return 1
        elif (id == 17): return 4
        elif (id == 18): return 4
        elif (id == 19): return getElectiveHours(4)
        elif (id == 20): return getElectiveHours(5)
        elif (id == 21): return 3
        elif (id == 22): return 3
        elif (id == 23): return 3
        elif (id == 24): return getElectiveHours(6)
        elif (id == 25): return getElectiveHours(7)
        elif (id == 26): return 3
        elif (id == 27): return 3
        elif (id == 28): return getElectiveHours(8)
        elif (id == 29): return getElectiveHours(9)
        elif (id == 30): return getElectiveHours(10)
        elif (id == 31): return getElectiveHours(11)
        elif (id == 32): return 3
        elif (id == 33): return getElectiveHours(12)
        elif (id == 34): return getElectiveHours(13)
        elif (id == 35): return getElectiveHours(14)
        elif (id == 36): return getElectiveHours(15)
        elif (id == 37): return 3

    #get flowchart information and populate document
    entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
    len = 0
    for item in entry:
        for thing in item['flowcharts']: # Get number of flowcharts an account has
            len += 1
        for i in range(len):
            if (item['flowcharts'][i]["name"] == name): # write flowchart information to document
                count = []
                data = []
                for j in range(0, 14):
                    count.append(0)
                    data.append([])
                for j in range(1,38): #determine size of table
                    cur = number_helper(item['flowcharts'][i][str(j)])
                    cur = int(cur)
                    count[cur] += 1
                    data[cur].append(j)
                numRows = 0 #number of semesters
                numCols = 0 #max number of classes in a semester
                largestCol = 0
                for j in range(1, 13):
                    if (count[j] > 0):
                        numRows += 1
                        if (count[j] > numCols):
                            numCols = count[j]
                            largestCol = j
                numCols += 1 #add one for semester name
                #add data to document
                if (numRows == 0): 
                    document.add_paragraph('This flowchart has no data.')
                else:
                    row = 0
                    line1 = "Flowchart Name: " + name
                    today = date.today()
                    line2 = "Modified on " + str(today)
                    document.add_paragraph(line1)
                    document.add_paragraph(line2)
                    document.add_paragraph("Note: If you have not specified all course names and hours for electives the total semester hours may not be accurate")
                    lineTaken = "Hours Taken: " + str(hoursTaken)
                    linePlanned = "Hours Planned: " + str(hoursPlanned)
                    document.add_paragraph(lineTaken)
                    document.add_paragraph(linePlanned)
                    if (largestCol != 1):
                        table = document.add_table(rows=numRows, cols=numCols+1)
                        table.style = 'TableGrid'
                        for j in range(1, 13):
                            if (count[j] > 0):
                                hoursRow = 0
                                curRow = table.rows[row].cells
                                paragraph = curRow[0].paragraphs[0]
                                run = paragraph.add_run(getSemesterName(j))
                                run.bold = True
                                for k in range(1, count[j]+1):
                                    curRow = table.rows[row].cells
                                    curRow[k].text = getClassName(data[j][k-1])
                                    hoursRow += getHours(data[j][k-1])
                                curRow[numCols].text = str(hoursRow) + " hours"
                                row += 1
                    else: 
                        newNumCols = 0
                        for j in range(1, 13):
                            if ((count[j] > newNumCols) and (j != 1)):
                                    newNumCols = count[j]
                        if (newNumCols == 0):
                            newNumCols = count[1]
                        rowIterations = count[1] / newNumCols
                        rowIterations = math.ceil(rowIterations)
                        newNumRows = numRows + rowIterations - 1
                        table = document.add_table(rows=newNumRows, cols=newNumCols+2)
                        table.style = 'TableGrid'
                        for j in range(1, 13):
                            hoursRow = 0
                            if (j != 1):
                                if (count[j] > 0):
                                    # print("adding ", getSemesterName(j), "to 0 in row ", row)
                                    curRow = table.rows[row].cells
                                    paragraph = curRow[0].paragraphs[0]
                                    run = paragraph.add_run(getSemesterName(j))
                                    run.bold = True
                                    for k in range(1, count[j]+1):
                                        curRow = table.rows[row].cells
                                        curRow[k].text = getClassName(data[j][k-1])
                                        hoursRow += getHours(data[j][k-1])
                                        # print("adding ", getClassName(data[j][k-1]), "to ", k, "in row ", row, "class ", data[j][k-1])
                                    curRow[newNumCols+1].text = str(hoursRow) + " hours"
                                    row += 1
                            else:
                                place = 0
                                if (count[j] > 0):
                                    for i in range(0, rowIterations):
                                        curRow = table.rows[row].cells
                                        if (i == 0):
                                            paragraph = curRow[0].paragraphs[0]
                                            run = paragraph.add_run(getSemesterName(j))
                                            run.bold = True
                                            # print("adding ", getSemesterName(j), "to 0 in row ", row)
                                            for k in range(1, newNumCols+1):
                                                curRow = table.rows[row].cells
                                                curRow[k].text = getClassName(data[j][place])
                                                hoursRow += getHours(data[j][k-1])
                                                # print("adding ", getClassName(data[j][place]), "to ", k, "in row ", row, "class ", data[j][place])
                                                place += 1
                                            row += 1
                                        else:
                                            for k in range(1, newNumCols+1):
                                                curRow = table.rows[row].cells
                                                curRow[k].text = getClassName(data[j][place])
                                                hoursRow += getHours(data[j][k-1])
                                                # print("adding ", getClassName(data[j][place]), "to ", k, "in row ", row, "class ", data[j][place])
                                                place += 1
                                                if (place >= count[1]):
                                                    break
                                            row += 1  
                                            if (row == rowIterations):
                                                curRow[newNumCols+1].text = str(hoursRow) + " hours"
    
    root_url = request.url_root
    ip_local = 'http://127.0.0.1:5000/'
    localhost = 'http://localhost:5000/'

    #send document to user to be downloaded 
    if ((root_url != ip_local) and (root_url != localhost)):
        saveName1 = "/tmp/" + title + ".docx"
    else:
        saveName1 = title + ".docx"
    document.save(saveName1)
    return saveName1

@app.route("/flowchart-new", methods=['post', 'get'])
def flowchart1():
    if request.method == "POST":
        email = session["email"]
        year = records.find({"email": email}, {"year": 1, "_id": 0})
        for item in year:
            yearPass = item['year']
        yearPass2 = int(yearPass)+1
        yearPass3 = int(yearPass)+2
        yearPass4 = int(yearPass)+3
        yearPass5 = int(yearPass)+4
        yearData = []
        yearData.append(int(yearPass))
        yearData.append(yearPass2)
        yearData.append(yearPass3)
        yearData.append(yearPass4)
        yearData.append(yearPass5)
        name = request.form.get("flowchartname")
        d1 = request.form.get("d1")
        d2 = request.form.get("d2")
        d3 = request.form.get("d3")
        d4 = request.form.get("d4")
        d5 = request.form.get("d5")
        d6 = request.form.get("d6")
        d7 = request.form.get("d7")
        d8 = request.form.get("d8")
        d9 = request.form.get("d9")
        d10 = request.form.get("d10")
        d11 = request.form.get("d11")
        d12 = request.form.get("d12")
        d13 = request.form.get("d13")
        d14 = request.form.get("d14")
        d15 = request.form.get("d15")
        d16 = request.form.get("d16")
        d17 = request.form.get("d17")
        d18 = request.form.get("d18")
        d19 = request.form.get("d19")
        d20 = request.form.get("d20")
        d21 = request.form.get("d21")
        d22 = request.form.get("d22")
        d23 = request.form.get("d23")
        d24 = request.form.get("d24")
        d25 = request.form.get("d25")
        d26 = request.form.get("d26")
        d27 = request.form.get("d27")
        d28 = request.form.get("d28")
        d29 = request.form.get("d29")
        d30 = request.form.get("d30")
        d31 = request.form.get("d31")
        d32 = request.form.get("d32")
        d33 = request.form.get("d33")
        d34 = request.form.get("d34")
        d35 = request.form.get("d35")
        d36 = request.form.get("d36")
        d37 = request.form.get("d37")
        
        # color information 
        c1 = request.form.get("c1")
        c2 = request.form.get("c2")
        c3 = request.form.get("c3")
        c4 = request.form.get("c4")
        c5 = request.form.get("c5")
        c6 = request.form.get("c6")
        c7 = request.form.get("c7")
        c8 = request.form.get("c8")
        c9 = request.form.get("c9")
        c10 = request.form.get("c10")
        c11 = request.form.get("c11")
        c12 = request.form.get("c12")
        c13 = request.form.get("c13")
        colArr = [c1, c2, c3, c4, c5, c6, c7, c8, c9, c10, c11, c12, c13]
        print(colArr)

        d7el = request.form.get("7electiveText")
        d7hrs = request.form.get("7hours")
        d10el = request.form.get("10electiveText")
        d10hrs = request.form.get("10hours")
        d14el = request.form.get("14electiveText")
        d14hrs = request.form.get("14hours")
        d15el = request.form.get("15electiveText")
        d15hrs = request.form.get("15hours")
        d19el = request.form.get("19electiveText")
        d19hrs = request.form.get("19hours")
        d20el = request.form.get("20electiveText")
        d20hrs = request.form.get("20hours")
        d24el = request.form.get("24electiveText")
        d24hrs = request.form.get("24hours")
        d25el = request.form.get("25electiveText")
        d25hrs = request.form.get("25hours")
        d28el = request.form.get("28electiveText")
        d28hrs = request.form.get("28hours")
        d29el = request.form.get("29electiveText")
        d29hrs = request.form.get("29hours")
        d30el = request.form.get("30electiveText")
        d30hrs = request.form.get("30hours")
        d31el = request.form.get("31electiveText")
        d31hrs = request.form.get("31hours")
        d33el = request.form.get("33electiveText")
        d33hrs = request.form.get("33hours")
        d34el = request.form.get("34electiveText")
        d34hrs = request.form.get("34hours")
        d35el = request.form.get("35electiveText")
        d35hrs = request.form.get("35hours")
        d36el = request.form.get("36electiveText")
        d36hrs = request.form.get("36hours")
        elNameArr = [d7el, d10el, d14el, d15el, d19el, d20el, d24el, d25el,
                     d28el, d29el, d30el, d31el, d33el, d34el, d35el, d36el]
        elHoursArr = [d7hrs, d10hrs, d14hrs, d15hrs, d19hrs, d20hrs, d24hrs, d25hrs,
                     d28hrs, d29hrs, d30hrs, d31hrs, d33hrs, d34hrs, d35hrs, d36hrs]
        print(f'\nNAME ARR:\n{elNameArr}\nHOURS ARRAY:\n{elHoursArr}\n')

        #checking if name is found
        email = session["email"]
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        len = 0
        name_found = False
        for item in entry:
            for thing in item['flowcharts']:
                len += 1
                for i in range(len):
                    if (item['flowcharts'][i]["name"] == name):
                        name_found = True
                        break

        data = [name, d1, d2, d3, d4, d5, d6, d7, d8, d9, d10, d11, d12, d13,
                d14, d15, d16, d17, d18, d19, d20, d21, d22, d23, d24, d25, 
                d26, d27, d28, d29, d30, d31, d32, d33, d34, d35, d36, d37]
        
        # appending color array
        data.append(colArr)

        # appending elective info
        data.append(elNameArr)
        data.append(elHoursArr)
        
        if (name == ""):
            message = 'Please enter a name for your flowchart'
            return render_template('flowchart-new.html', message=message, data=data, yearData=yearData)
        elif (name_found):
            message = 'You already have a flowchart named ' + name
            return render_template('flowchart-new.html', message=message, data=data, yearData=yearData)
        else: 
            email = session["email"]
            len = 0
            entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
            for item in entry:
                for thing in item['flowcharts']:
                    len += 1
            records.update_one(
                {"email": email},
                {"$push": {
                    "flowcharts": {
                        "$each": [{ "name": name, 
                            "1": d1, "2": d2, "3": d3, "4": d4, "5": d5,
                            "6": d6, "7": d7, "8": d8, "9": d9, 
                            "10": d10, "11": d11, "12": d12, "13": d13,
                            "14": d14, "15": d15, "16": d16, "17": d17, "18": d18,
                            "19": d19, "20": d20, "21": d21, "22": d22, "23": d23,
                            "24": d24, "25": d25, "26": d26, "27": d27, "28": d28,
                            "29": d29, "30": d30, "31": d31, "32": d32, "33": d33,
                            "34": d34, "35": d35, "36": d36, "37": d37, "colors": colArr,
                            "elective_names": elNameArr, "elective_hours": elHoursArr}],
                        "$position": len
                    }
                }
            })   
            message = 'Your flowchart has been saved'
            return redirect(url_for('flowchart2', name=name))#render_template('flowchart-edit.html', message=message, name=name)
    if "email" in session: # GET
        email = session["email"]
        year = records.find({"email": email}, {"year": 1, "_id": 0})
        for item in year:
            yearPass = item['year']
        yearPass2 = int(yearPass)+1
        yearPass3 = int(yearPass)+2
        yearPass4 = int(yearPass)+3
        yearPass5 = int(yearPass)+4
        yearData = []
        yearData.append(int(yearPass))
        yearData.append(yearPass2)
        yearData.append(yearPass3)
        yearData.append(yearPass4)
        yearData.append(yearPass5)
        return render_template('flowchart-new.html', yearData = yearData)
    else:
        return redirect(url_for("login"))

@app.route("/flowchart-edit/<name>", methods=['post', 'get'])
def flowchart2(name):
    if request.method == "POST":    # Get data from html page
        print("name: ", name)
        d1 = request.form.get("d1")
        d2 = request.form.get("d2")
        d3 = request.form.get("d3")
        d4 = request.form.get("d4")
        d5 = request.form.get("d5")
        d6 = request.form.get("d6")
        d7 = request.form.get("d7")
        d8 = request.form.get("d8")
        d9 = request.form.get("d9")
        d10 = request.form.get("d10")
        d11 = request.form.get("d11")
        d12 = request.form.get("d12")
        d13 = request.form.get("d13")
        d14 = request.form.get("d14")
        d15 = request.form.get("d15")
        d16 = request.form.get("d16")
        d17 = request.form.get("d17")
        d18 = request.form.get("d18")
        d19 = request.form.get("d19")
        d20 = request.form.get("d20")
        d21 = request.form.get("d21")
        d22 = request.form.get("d22")
        d23 = request.form.get("d23")
        d24 = request.form.get("d24")
        d25 = request.form.get("d25")
        d26 = request.form.get("d26")
        d27 = request.form.get("d27")
        d28 = request.form.get("d28")
        d29 = request.form.get("d29")
        d30 = request.form.get("d30")
        d31 = request.form.get("d31")
        d32 = request.form.get("d32")
        d33 = request.form.get("d33")
        d34 = request.form.get("d34")
        d35 = request.form.get("d35")
        d36 = request.form.get("d36")
        d37 = request.form.get("d37")
        
        # color information 
        c1 = request.form.get("c1")
        c2 = request.form.get("c2")
        c3 = request.form.get("c3")
        c4 = request.form.get("c4")
        c5 = request.form.get("c5")
        c6 = request.form.get("c6")
        c7 = request.form.get("c7")
        c8 = request.form.get("c8")
        c9 = request.form.get("c9")
        c10 = request.form.get("c10")
        c11 = request.form.get("c11")
        c12 = request.form.get("c12")
        c13 = request.form.get("c13")
        colArr = [c1, c2, c3, c4, c5, c6, c7, c8, c9, c10, c11, c12, c13]
        print(colArr)


        """ Check elective info. - Render template with elective info from DB """
        # get the name and hours arrays from the database 
        email = session["email"]
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        length = 0
        for item in entry:
            for thing in item['flowcharts']: # Get number of flowcharts an account has
                length += 1
            for i in range(length):
                if (item['flowcharts'][i]["name"] == name): # Get a specific flowchart using the name to find it; set the data
                    testName = item['flowcharts'][i]["elective_names"]
                    testHours = item['flowcharts'][i]["elective_hours"]
                
        print(f'\t**Arrays from DB\n\ttestName: {testName}\n\ttestHours: {testHours}\n')

        # check to see what elective information is already added
        elective_ids = [7,10,14,15,19,20,24,25,28,29,30,31,33,34,35,36]
        elNameArr = []
        elHoursArr = []
        for (id, temp_name, temp_hour) in zip(elective_ids, testName, testHours):
            id = str(id)
            formName = request.form.get(str(id+"electiveText"))
            formHour = request.form.get(str(id+"hours"))

            if temp_name == '':
                # There is NOT information in the DB                
                ### For troubleshooting - uncomment following 2 lines ###
                # print(f'\t** name id: {str(id+"electiveText")}\thour id: {str(id+"hours")} **')
                # print(f'\t** formName: {formName}\tformHour: {formHour} **')
                elNameArr.append(formName)
                elHoursArr.append(formHour)
            else:
                # There is information in the DB #
                # Check if there is conflicting info submitted to the form
                if(formName != '' and formName != temp_name):
                    # Overwrite DB with new info
                    elNameArr.append(formName)
                    elHoursArr.append(formHour)                    
                else:
                    # Get info from DB
                    elNameArr.append(temp_name)
                    elHoursArr.append(temp_hour)

        print(f'\nNAME ARR:\n{elNameArr}\nHOURS ARRAY:\n{elHoursArr}\n')

        email = session["email"]
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        newData = []
        len = 0
        for item in entry:
            for thing in item['flowcharts']: # Get number of flowcharts an account has
                len += 1
            for i in range(len):
                if (item['flowcharts'][i]["name"] == name): # Get a specific flowchart using the name to find it; set the data
                    item['flowcharts'][i]["1"] = d1
                    item['flowcharts'][i]["2"] = d2
                    item['flowcharts'][i]["3"] = d3
                    item['flowcharts'][i]["4"] = d4
                    item['flowcharts'][i]["5"] = d5
                    item['flowcharts'][i]["6"] = d6
                    item['flowcharts'][i]["7"] = d7
                    item['flowcharts'][i]["8"] = d8
                    item['flowcharts'][i]["9"] = d9
                    item['flowcharts'][i]["10"] = d10
                    item['flowcharts'][i]["11"] = d11
                    item['flowcharts'][i]["12"] = d12
                    item['flowcharts'][i]["13"] = d13
                    item['flowcharts'][i]["14"] = d14
                    item['flowcharts'][i]["15"] = d15
                    item['flowcharts'][i]["16"] = d16
                    item['flowcharts'][i]["17"] = d17
                    item['flowcharts'][i]["18"] = d18
                    item['flowcharts'][i]["19"] = d19
                    item['flowcharts'][i]["20"] = d20
                    item['flowcharts'][i]["21"] = d21
                    item['flowcharts'][i]["22"] = d22
                    item['flowcharts'][i]["23"] = d23
                    item['flowcharts'][i]["24"] = d24
                    item['flowcharts'][i]["25"] = d25
                    item['flowcharts'][i]["26"] = d26
                    item['flowcharts'][i]["27"] = d27
                    item['flowcharts'][i]["28"] = d28
                    item['flowcharts'][i]["29"] = d29
                    item['flowcharts'][i]["30"] = d30
                    item['flowcharts'][i]["31"] = d31
                    item['flowcharts'][i]["32"] = d32
                    item['flowcharts'][i]["33"] = d33
                    item['flowcharts'][i]["34"] = d34
                    item['flowcharts'][i]["35"] = d35
                    item['flowcharts'][i]["36"] = d36
                    item['flowcharts'][i]["37"] = d37
                    
                    # color info to database
                    item['flowcharts'][i]["colors"] = colArr

                    # elective information
                    item['flowcharts'][i]["elective_names"] = elNameArr
                    item['flowcharts'][i]["elective_hours"] = elHoursArr

                newData.append(item['flowcharts'][i])
        
        data = [name, d1, d2, d3, d4, d5, d6, d7, d8, d9, d10, d11, d12, d13,
                d14, d15, d16, d17, d18, d19, d20, d21, d22, d23, d24, d25, 
                d26, d27, d28, d29, d30, d31, d32, d33, d34, d35, d36, d37]
        
        # appending color array - at index 38
        data.append(colArr)

        # appending elective info - at index 39 & 40
        data.append(elNameArr)
        data.append(elHoursArr)

        print("data: ", data)
        records.update_one(
            {"email": email},
            {"$set": {"flowcharts": newData}}
        )
        message = 'Your flowchart has been saved'
        year = records.find({"email": email}, {"year": 1, "_id": 0})
        for item in year:
            yearPass = item['year']
        yearPass2 = int(yearPass)+1
        yearPass3 = int(yearPass)+2
        yearPass4 = int(yearPass)+3
        yearPass5 = int(yearPass)+4
        yearData = []
        yearData.append(int(yearPass))
        yearData.append(yearPass2)
        yearData.append(yearPass3)
        yearData.append(yearPass4)
        yearData.append(yearPass5)

        export = request.form.get("export")
        if (export == "on"):
            print("EXPORT")
            hoursTaken = request.form.get("hoursTaken")
            hoursPlanned = request.form.get("hoursPlanned")
            years2 = []
            for i in range(0, 5):
                years2.append(yearData[i])
            print("years:",  years2)
            fileName = exportData(years2, email, name, elNameArr, elHoursArr)
            return send_file(fileName, name, as_attachment=True, download_name=fileName)

        return render_template('flowchart-edit.html', message=message, name=name, data=data, yearData=yearData)
    # GET METHOD
    if "email" in session: # GET
        print("edit name: ", name)
        email = session["email"]
        entry = records.find({"email": email}, {"flowcharts": 1, "_id": 0})
        data = []
        len = 0
        for item in entry:
            for thing in item['flowcharts']:
                len += 1
            for i in range(len):
                if (item['flowcharts'][i]["name"] == name):
                    data.append(name)
                    data.append(item['flowcharts'][i]["1"])
                    data.append(item['flowcharts'][i]["2"])
                    data.append(item['flowcharts'][i]["3"])
                    data.append(item['flowcharts'][i]["4"])
                    data.append(item['flowcharts'][i]["5"])
                    data.append(item['flowcharts'][i]["6"])
                    data.append(item['flowcharts'][i]["7"])
                    data.append(item['flowcharts'][i]["8"])
                    data.append(item['flowcharts'][i]["9"])
                    data.append(item['flowcharts'][i]["10"])
                    data.append(item['flowcharts'][i]["11"])
                    data.append(item['flowcharts'][i]["12"])
                    data.append(item['flowcharts'][i]["13"])
                    data.append(item['flowcharts'][i]["14"])
                    data.append(item['flowcharts'][i]["15"])
                    data.append(item['flowcharts'][i]["16"])
                    data.append(item['flowcharts'][i]["17"])
                    data.append(item['flowcharts'][i]["18"])
                    data.append(item['flowcharts'][i]["19"])
                    data.append(item['flowcharts'][i]["20"])
                    data.append(item['flowcharts'][i]["21"])
                    data.append(item['flowcharts'][i]["22"])
                    data.append(item['flowcharts'][i]["23"])
                    data.append(item['flowcharts'][i]["24"])
                    data.append(item['flowcharts'][i]["25"])
                    data.append(item['flowcharts'][i]["26"])
                    data.append(item['flowcharts'][i]["27"])
                    data.append(item['flowcharts'][i]["28"])
                    data.append(item['flowcharts'][i]["29"])
                    data.append(item['flowcharts'][i]["30"])
                    data.append(item['flowcharts'][i]["31"])
                    data.append(item['flowcharts'][i]["32"])
                    data.append(item['flowcharts'][i]["33"])
                    data.append(item['flowcharts'][i]["34"])
                    data.append(item['flowcharts'][i]["35"])
                    data.append(item['flowcharts'][i]["36"])
                    data.append(item['flowcharts'][i]["37"])
                    
                    # color information - from dictionary
                    data.append(item['flowcharts'][i]["colors"])

                    # elective information - from dictionary
                    data.append(item['flowcharts'][i]["elective_names"])
                    data.append(item['flowcharts'][i]["elective_hours"])

        year = records.find({"email": email}, {"year": 1, "_id": 0})
        for item in year:
            yearPass = item['year']
        yearPass2 = int(yearPass)+1
        yearPass3 = int(yearPass)+2
        yearPass4 = int(yearPass)+3
        yearPass5 = int(yearPass)+4
        yearData = []
        yearData.append(int(yearPass))
        yearData.append(yearPass2)
        yearData.append(yearPass3)
        yearData.append(yearPass4)
        yearData.append(yearPass5)
        print("GET data: ", data)
        return render_template('flowchart-edit.html', data=data, yearData=yearData)
    else:
        return redirect(url_for("login"))


# FIXME: is this route being used?
@app.route("/edit_elective", methods=["POST", "GET"])
def edit_elective():
    if request.method == "POST":
        if request.form.get("7Submit"):
            info_dict = {
                'show_submit': True,
                'show_edit': False
            }
            return render_template('flowchart-new.html', info_dict=info_dict)
        elif request.form.get("7Edit"):
            info_dict = {
                'show_submit': False,
                'show_edit': True
            }
            return render_template('flowchart-new.html', info_dict=info_dict)


#end of code to run it
if __name__ == "__main__":
  app.run(debug=True)